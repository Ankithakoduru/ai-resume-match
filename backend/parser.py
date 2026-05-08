import io
import re
from typing import Optional

import pdfplumber
from docx import Document


SKILL_KEYWORDS = [
    "Python", "SQL", "Java", "JavaScript", "TypeScript", "R", "Scala", "Go",
    "Power BI", "Tableau", "Looker", "Qlik", "Metabase", "Excel", "DAX",
    "Power Query", "SSRS", "SSIS", "Alteryx", "SAS",
    "AWS", "Azure", "GCP", "Snowflake", "Databricks", "Spark", "Airflow",
    "dbt", "Redshift", "BigQuery", "Synapse",
    "Pandas", "NumPy", "Scikit-learn", "TensorFlow", "PyTorch", "Keras",
    "Matplotlib", "Seaborn", "Plotly", "Jupyter",
    "PostgreSQL", "MySQL", "MongoDB", "SQLite", "Oracle", "MSSQL",
    "React", "Vue", "Angular", "Node.js", "FastAPI", "Django", "Flask",
    "Docker", "Kubernetes", "Git", "GitHub", "Jira", "Confluence",
    "SAP", "Salesforce", "HubSpot", "Workday",
]


def extract_text_from_pdf(content: bytes) -> str:
    text = ""
    try:
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        print(f"PDF parsing error: {e}")
    return text


def extract_text_from_docx(content: bytes) -> str:
    text = ""
    try:
        doc = Document(io.BytesIO(content))
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + "\n"
    except Exception as e:
        print(f"DOCX parsing error: {e}")
    return text


def extract_name(text: str) -> str:
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    if lines:
        first_line = lines[0]
        if len(first_line.split()) <= 5 and not any(c.isdigit() for c in first_line):
            return first_line
    return "Unknown Candidate"


def extract_email(text: str) -> Optional[str]:
    match = re.search(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", text)
    return match.group(0) if match else None


def extract_phone(text: str) -> Optional[str]:
    match = re.search(r"(\+?1?\s?)?(\(?\d{3}\)?[\s\-.]?\d{3}[\s\-.]?\d{4})", text)
    return match.group(0).strip() if match else None


def extract_location(text: str) -> Optional[str]:
    patterns = [
        r"\b([A-Z][a-z]+(?:\s[A-Z][a-z]+)?),\s*([A-Z]{2})\b",
        r"\b([A-Z][a-z]+(?:\s[A-Z][a-z]+)?),\s*(Florida|California|Texas|New York|Georgia|Illinois)\b",
    ]
    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            return match.group(0)
    return None


def extract_skills(text: str) -> list:
    found = []
    text_lower = text.lower()
    for skill in SKILL_KEYWORDS:
        if skill.lower() in text_lower:
            found.append(skill)
    return found


def extract_experience_years(text: str) -> Optional[int]:
    patterns = [
        r"(\d+)\+?\s*years?\s*(?:of\s*)?experience",
        r"experience[:\s]+(\d+)\+?\s*years?",
    ]
    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            return int(match.group(1))
    year_matches = re.findall(r"\b(19|20)\d{2}\b", text)
    if len(year_matches) >= 2:
        years = sorted([int(y) for y in year_matches])
        return max(1, years[-1] - years[0])
    return None


def extract_education(text: str) -> Optional[str]:
    degrees = ["Ph.D", "PhD", "M.S.", "M.Sc", "M.B.A", "MBA", "B.S.", "B.Sc", "B.A.", "Bachelor", "Master", "Doctorate"]
    for degree in degrees:
        pattern = rf"{re.escape(degree)}[\w\s.,]{{0,60}}"
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            return match.group(0).strip()[:80]
    return None


def extract_salary(text: str) -> Optional[str]:
    match = re.search(r"\$\s*(\d{2,3}(?:,\d{3})?(?:\s*[kK])?)\s*(?:/yr|/year|annually|per year)?", text)
    if match:
        return match.group(0).strip()
    return None


def parse_resume(content: bytes, filename: str) -> dict:
    if filename.lower().endswith(".pdf"):
        text = extract_text_from_pdf(content)
    else:
        text = extract_text_from_docx(content)

    if not text.strip():
        return {
            "name": filename.replace(".pdf", "").replace(".docx", ""),
            "raw_text": "",
            "skills": [],
            "experience_years": None,
            "education": None,
            "location": None,
            "email": None,
            "phone": None,
            "salary": None,
        }

    return {
        "name": extract_name(text),
        "raw_text": text,
        "skills": extract_skills(text),
        "experience_years": extract_experience_years(text),
        "education": extract_education(text),
        "location": extract_location(text),
        "email": extract_email(text),
        "phone": extract_phone(text),
        "salary": extract_salary(text),
    }
